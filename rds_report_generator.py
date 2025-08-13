import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime
import numpy as np
import warnings
warnings.filterwarnings('ignore')

class RDSReportGeneratorWord:
    def __init__(self, base_output_dir='rds_report', template_path=None, cost_optimization_csv=None):
        self.base_output_dir = base_output_dir
        self.template_path = template_path
        self.cost_optimization_csv = cost_optimization_csv
        self.metrics_data = {}
        self.instance_names = []
        self.dates = None
        self.cost_optimization_data = None
        
        # Configurar estilo dos gráficos
        plt.style.use('default')
        sns.set_palette("husl")
        
        # Criar diretórios
        os.makedirs(base_output_dir, exist_ok=True)
        os.makedirs(f"{base_output_dir}/charts", exist_ok=True)
        
        print(f"Diretórios criados: {base_output_dir}")
    
    def load_cost_optimization_data(self):
        """Carrega dados de otimização de custos"""
        if not self.cost_optimization_csv or not os.path.exists(self.cost_optimization_csv):
            print("⚠️ Arquivo de otimização de custos não encontrado ou não fornecido")
            return
        
        try:
            print(f"Carregando dados de otimização de custos: {self.cost_optimization_csv}")
            
            # Ler CSV de otimização
            df = pd.read_csv(self.cost_optimization_csv)
            
            # Criar um dicionário indexado por nome da instância para busca rápida
            self.cost_optimization_data = {}
            
            for _, row in df.iterrows():
                instance_id = row.get('DB_Instance_ID', '')
                instance_name = row.get('DB_Instance_Name', '')
                
                # Usar ID como chave principal, nome como secundária
                key = instance_id if instance_id else instance_name
                
                if key:
                    self.cost_optimization_data[key] = {
                        'instance_name': instance_name,
                        'instance_id': instance_id,
                        'status': row.get('Status', 'unknown'),
                        'engine': row.get('Engine', 'unknown'),
                        'current_instance_class': row.get('Current_Instance_Class', 'unknown'),
                        'finding': row.get('Finding', 'unknown'),
                        'finding_reason_codes': row.get('Finding_Reason_Codes', ''),  # NOVA INFORMAÇÃO
                        'recommended_instance_class': row.get('Recommended_Instance_Class', ''),
                        'current_ondemand_monthly': self.parse_currency(row.get('Current_OnDemand_Monthly', 0)),
                        'recommended_ondemand_monthly': self.parse_currency(row.get('Recommended_OnDemand_Monthly', 0)),
                        'ondemand_monthly_savings': self.parse_currency(row.get('OnDemand_Monthly_Savings', 0)),
                        'ondemand_savings_percentage': self.parse_percentage(row.get('OnDemand_Savings_Percentage', 0)),
                        'current_cpu_util': self.parse_percentage(row.get('Current_Max_CPU_Util', 0)),
                        'current_memory_util': self.parse_percentage(row.get('Current_Max_Memory_Util', 0)),
                        'recommended_cpu_util': self.parse_percentage(row.get('Recommended_Max_CPU_Util', 0)),
                        # Modalidades de reserva ATUAIS
                        'current_no_upfront_monthly_1yr': self.parse_currency(row.get('Current_No_Upfront_Monthly_1yr', 0)),
                        'current_partial_upfront_initial_1yr': self.parse_currency(row.get('Current_Partial_Upfront_Initial_1yr', 0)),
                        'current_partial_upfront_monthly_1yr': self.parse_currency(row.get('Current_Partial_Upfront_Monthly_1yr', 0)),
                        'current_all_upfront_1yr': self.parse_currency(row.get('Current_All_Upfront_1yr', 0)),
                        'current_no_upfront_monthly_3yr': self.parse_currency(row.get('Current_No_Upfront_Monthly_3yr', 0)),
                        'current_partial_upfront_initial_3yr': self.parse_currency(row.get('Current_Partial_Upfront_Initial_3yr', 0)),
                        'current_partial_upfront_monthly_3yr': self.parse_currency(row.get('Current_Partial_Upfront_Monthly_3yr', 0)),
                        'current_all_upfront_3yr': self.parse_currency(row.get('Current_All_Upfront_3yr', 0)),
                        # Modalidades RECOMENDADAS
                        'recommended_no_upfront_monthly_1yr': self.parse_currency(row.get('Recommended_No_Upfront_Monthly_1yr', 0)),
                        'recommended_partial_upfront_initial_1yr': self.parse_currency(row.get('Recommended_Partial_Upfront_Initial_1yr', 0)),
                        'recommended_partial_upfront_monthly_1yr': self.parse_currency(row.get('Recommended_Partial_Upfront_Monthly_1yr', 0)),
                        'recommended_all_upfront_1yr': self.parse_currency(row.get('Recommended_All_Upfront_1yr', 0)),
                        'recommended_no_upfront_monthly_3yr': self.parse_currency(row.get('Recommended_No_Upfront_Monthly_3yr', 0)),
                        'recommended_partial_upfront_initial_3yr': self.parse_currency(row.get('Recommended_Partial_Upfront_Initial_3yr', 0)),
                        'recommended_partial_upfront_monthly_3yr': self.parse_currency(row.get('Recommended_Partial_Upfront_Monthly_3yr', 0)),
                        'recommended_all_upfront_3yr': self.parse_currency(row.get('Recommended_All_Upfront_3yr', 0)),
                        # Savings
                        'no_upfront_1yr_monthly_savings': self.parse_currency(row.get('No_Upfront_1yr_Monthly_Savings', 0)),
                        'all_upfront_1yr_annual_savings': self.parse_currency(row.get('All_Upfront_1yr_Annual_Savings', 0)),
                        'all_upfront_3yr_annual_savings': self.parse_currency(row.get('All_Upfront_3yr_Annual_Savings', 0))
                    }
            
            print(f"✓ Dados de otimização carregados: {len(self.cost_optimization_data)} instâncias")
            
        except Exception as e:
            print(f"✗ Erro ao carregar dados de otimização: {str(e)}")
    
    def parse_currency(self, value):
        """Converte string de moeda para float"""
        if pd.isna(value) or value == '' or value == 'not available':
            return 0.0
        if isinstance(value, str):
            # Remove $, espaços e converte para float
            cleaned = value.replace('$', '').replace(',', '').replace(' ', '')
            try:
                return float(cleaned)
            except:
                return 0.0
        return float(value) if value else 0.0
    
    def parse_percentage(self, value):
        """Converte string de percentual para float"""
        if pd.isna(value) or value == '' or value == 'not available':
            return 0.0
        if isinstance(value, str):
            # Remove % e converte para float
            cleaned = value.replace('%', '').replace(' ', '')
            try:
                return float(cleaned)
            except:
                return 0.0
        return float(value) if value else 0.0
    
    def format_finding_reason_codes(self, reason_codes):
        """Formata e traduz os códigos de razão do finding"""
        if not reason_codes or reason_codes == 'not available':
            return ""
        
        # Dicionário de tradução dos códigos
        reason_translations = {
            'CPUUnderprovisioned': 'CPU sub-provisionada',
            'CPUOverprovisioned': 'CPU super-provisionada',
            'MemoryUnderprovisioned': 'Memória sub-provisionada',
            'MemoryOverprovisioned': 'Memória super-provisionada',
            'EBSIOPSUnderprovisioned': 'IOPS EBS sub-provisionado',
            'EBSIOPSOverprovisioned': 'IOPS EBS super-provisionado',
            'EBSThroughputUnderprovisioned': 'Throughput EBS sub-provisionado',
            'EBSThroughputOverprovisioned': 'Throughput EBS super-provisionado',
            'NetworkBandwidthUnderprovisioned': 'Largura de banda de rede sub-provisionada',
            'NetworkBandwidthOverprovisioned': 'Largura de banda de rede super-provisionada',
            'NewEngineVersionAvailable': 'Nova versão do engine disponível'
        }
        
        # Separar códigos por vírgula e traduzir
        codes = reason_codes.split(',')
        translated_codes = []
        
        for code in codes:
            code = code.strip()
            if code in reason_translations:
                translated_codes.append(reason_translations[code])
            else:
                translated_codes.append(code)  # Manter original se não encontrar tradução
        
        return ', '.join(translated_codes)
    
    def get_cost_optimization_info(self, instance_name):
        """Busca informações de otimização para uma instância"""
        # Tentar busca exata primeiro
        if instance_name in self.cost_optimization_data:
            return self.cost_optimization_data[instance_name]
        
        # Busca parcial (caso os nomes não batam exatamente)
        for key, data in self.cost_optimization_data.items():
            if (instance_name.lower() in key.lower() or 
                key.lower() in instance_name.lower() or
                instance_name.lower() in data['instance_name'].lower() or
                data['instance_name'].lower() in instance_name.lower()):
                return data
        
        return None
    
    def get_reservation_options(self, opt_info, use_recommended=False):
        """Retorna todas as opções de reserva disponíveis"""
        if not opt_info:
            return []
        
        options = []
        prefix = 'recommended' if use_recommended else 'current'
        
        # 1 ano - No Upfront
        no_upfront_1yr = opt_info.get(f'{prefix}_no_upfront_monthly_1yr', 0)
        if no_upfront_1yr > 0:
            annual_cost = no_upfront_1yr * 12
            options.append({
                'term': '1 ano',
                'type': 'No Upfront',
                'monthly_cost': no_upfront_1yr,
                'annual_cost': annual_cost,
                'upfront_cost': 0,
                'total_cost': annual_cost,
                'display': f"${no_upfront_1yr:.2f}/mês (${annual_cost:.0f}/ano)"
            })
        
        # 1 ano - All Upfront
        all_upfront_1yr = opt_info.get(f'{prefix}_all_upfront_1yr', 0)
        if all_upfront_1yr > 0:
            options.append({
                'term': '1 ano',
                'type': 'All Upfront',
                'monthly_cost': 0,
                'annual_cost': all_upfront_1yr,
                'upfront_cost': all_upfront_1yr,
                'total_cost': all_upfront_1yr,
                'display': f"${all_upfront_1yr:.0f} pagamento único"
            })
        
        # 3 anos - No Upfront
        no_upfront_3yr = opt_info.get(f'{prefix}_no_upfront_monthly_3yr', 0)
        if no_upfront_3yr > 0:
            annual_cost = no_upfront_3yr * 12
            total_cost = annual_cost * 3
            options.append({
                'term': '3 anos',
                'type': 'No Upfront',
                'monthly_cost': no_upfront_3yr,
                'annual_cost': annual_cost,
                'upfront_cost': 0,
                'total_cost': total_cost,
                'display': f"${no_upfront_3yr:.2f}/mês (${annual_cost:.0f}/ano)"
            })
        
        # 3 anos - All Upfront
        all_upfront_3yr = opt_info.get(f'{prefix}_all_upfront_3yr', 0)
        if all_upfront_3yr > 0:
            annual_cost = all_upfront_3yr / 3
            options.append({
                'term': '3 anos',
                'type': 'All Upfront',
                'monthly_cost': 0,
                'annual_cost': annual_cost,
                'upfront_cost': all_upfront_3yr,
                'total_cost': all_upfront_3yr,
                'display': f"${all_upfront_3yr:.0f} pagamento único (${annual_cost:.0f}/ano)"
            })
        
        return options
    
    def get_best_reservation_option(self, opt_info, use_recommended=False):
        """Determina a melhor opção de reserva baseada no menor custo anual"""
        options = self.get_reservation_options(opt_info, use_recommended)
        if not options:
            return None
        
        # Retornar a opção com menor custo anual
        best_option = min(options, key=lambda x: x['annual_cost'])
        return best_option

    def load_metric_data(self, csv_files):
        """
        Carrega dados de múltiplas métricas com estruturas diferentes
        """
        for metric_name, csv_path in csv_files.items():
            if os.path.exists(csv_path):
                print(f"Carregando dados de {metric_name} de {csv_path}...")
                try:
                    # Ler o arquivo linha por linha para identificar a estrutura
                    with open(csv_path, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                    
                    # Identificar linhas importantes
                    label_line_idx = None
                    data_start_line = 0
                    
                    for i, line in enumerate(lines):
                        if line.startswith('Label,'):
                            label_line_idx = i
                            data_start_line = i + 1
                            break
                        elif line.startswith('2025/') or line.startswith('20'):  # Linha de dados
                            data_start_line = i
                            break
                    
                    # Ler dados a partir da linha identificada
                    df = pd.read_csv(csv_path, skiprows=data_start_line)
                    
                    # Primeira vez - extrair datas e nomes das instâncias
                    if self.dates is None:
                        # Primeira coluna contém as datas
                        date_column = df.columns[0]
                        self.dates = pd.to_datetime(df[date_column])
                        
                        # Extrair nomes das instâncias da linha Label
                        if label_line_idx is not None:
                            label_line = lines[label_line_idx]
                            self.instance_names = [name.strip() for name in label_line.split(',')[1:]]
                        else:
                            # Usar nomes das colunas se não houver linha Label
                            self.instance_names = list(df.columns[1:])
                    
                    # Verificar se temos dados válidos nas colunas
                    numeric_data = df.iloc[:, 1:].copy()
                    
                    # Converter para numérico, colocando NaN onde não conseguir
                    for col in numeric_data.columns:
                        numeric_data[col] = pd.to_numeric(numeric_data[col], errors='coerce')
                    
                    # Para Free Storage Space, converter de bytes para GB
                    if 'storage' in metric_name.lower() or 'freestoragespace' in csv_path.lower():
                        print(f"Convertendo {metric_name} de bytes para GB...")
                        numeric_data = numeric_data / (1024**3)  # Converter bytes para GB
                        
                    # Para Freeable Memory, converter de bytes para GB  
                    elif 'memory' in metric_name.lower() or 'freeablememory' in csv_path.lower():
                        print(f"Convertendo {metric_name} de bytes para GB...")
                        numeric_data = numeric_data / (1024**3)  # Converter bytes para GB
                    
                    # Armazenar dados processados
                    self.metrics_data[metric_name] = numeric_data
                    print(f"✓ {metric_name}: {len(df)} registros carregados")
                    
                    # Debug: Mostrar sample dos dados
                    print(f"   Sample de dados para {metric_name}:")
                    for col_idx in range(min(3, len(numeric_data.columns))):
                        col_name = numeric_data.columns[col_idx]
                        sample_values = numeric_data[col_name].dropna().head(3).values
                        print(f"     {col_name}: {sample_values}")
                    
                except Exception as e:
                    print(f"✗ Erro ao carregar {csv_path}: {str(e)}")
            else:
                print(f"✗ Arquivo não encontrado: {csv_path}")
        
        print(f"Total de métricas carregadas: {len(self.metrics_data)}")
        print(f"Total de instâncias: {len(self.instance_names)}")
    
    def create_individual_charts(self):
        """Cria gráficos individuais para cada instância e métrica"""
        chart_paths = {}
        
        # Configuração das métricas com unidades corretas
        metrics_config = {
            'CPU': {'color': '#FF6B6B', 'ylabel': 'CPU Utilization (%)', 'ylim': (0, 100)},
            'Memory': {'color': '#4ECDC4', 'ylabel': 'Free Memory (GB)', 'ylim': None},
            'FreeableMemory': {'color': '#4ECDC4', 'ylabel': 'Free Memory (GB)', 'ylim': None},
            'Storage': {'color': '#96CEB4', 'ylabel': 'Free Storage (GB)', 'ylim': None},
            'FreeStorageSpace': {'color': '#96CEB4', 'ylabel': 'Free Storage (GB)', 'ylim': None},
            'DatabaseConnections': {'color': '#45B7D1', 'ylabel': 'Connections', 'ylim': None},
        }
        
        print(f"\nGerando gráficos para {len(self.instance_names)} instâncias...")
        
        for i, instance in enumerate(self.instance_names):
            if i % 10 == 0:
                print(f"Processando instância {i+1}/{len(self.instance_names)}: {instance}")
            
            try:
                # Determinar quantas métricas temos disponíveis
                available_metrics = list(self.metrics_data.keys())
                n_metrics = len(available_metrics)
                
                if n_metrics == 0:
                    continue
                
                # Configurar subplot baseado no número de métricas - otimizado para Word
                if n_metrics == 1:
                    fig, axes = plt.subplots(1, 1, figsize=(10, 6))
                    axes = [axes]
                elif n_metrics == 2:
                    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
                elif n_metrics <= 4:
                    fig, axes = plt.subplots(2, 2, figsize=(14, 11))
                    axes = axes.flatten()
                else:
                    # Para mais de 4 métricas, usar layout 3x2
                    fig, axes = plt.subplots(3, 2, figsize=(14, 16))
                    axes = axes.flatten()
                
                # Obter informações de otimização se disponível
                opt_info = self.get_cost_optimization_info(instance) if self.cost_optimization_data else None
                
                # Título com informações de otimização (sem emojis)
                title = f'Métricas de Performance - {instance}'
                if opt_info:
                    finding = opt_info['finding']
                    if finding == 'Optimized':
                        title += ' (OTIMIZADA)'
                    elif finding == 'Underprovisioned':
                        title += ' (SUB-PROVISIONADA)'
                    elif finding == 'Overprovisioned':
                        title += ' (SUPER-PROVISIONADA)'
                    else:
                        title += f' ({finding.upper()})'
                
                fig.suptitle(title, fontsize=13, fontweight='bold')
                
                for idx, metric_name in enumerate(available_metrics):
                    if idx >= len(axes):
                        break
                        
                    ax = axes[idx]
                    
                    # Buscar configuração da métrica (com fallback)
                    config = None
                    for key, cfg in metrics_config.items():
                        if key.lower() in metric_name.lower():
                            config = cfg
                            break
                    
                    if config is None:
                        config = {'color': '#333333', 'ylabel': metric_name, 'ylim': None}
                    
                    try:
                        # Verificar se temos dados para esta instância
                        if i >= len(self.metrics_data[metric_name].columns):
                            ax.text(0.5, 0.5, f'Instância não encontrada\nno dataset {metric_name}', 
                                   transform=ax.transAxes, ha='center', va='center')
                            ax.set_title(f'{metric_name}', fontsize=11, fontweight='bold')
                            continue
                        
                        # Obter dados da métrica para esta instância
                        metric_values = self.metrics_data[metric_name].iloc[:, i]
                        
                        # Verificar se há dados válidos
                        valid_data = metric_values.dropna()
                        
                        if len(valid_data) > 0:
                            # Plotar apenas onde temos dados válidos
                            valid_dates = self.dates[metric_values.notna()]
                            
                            ax.plot(valid_dates, valid_data, 
                                   color=config['color'], linewidth=2, 
                                   marker='o', markersize=2, alpha=0.8)
                            
                            # Configurar gráfico
                            ax.set_title(f'{metric_name}', fontsize=11, fontweight='bold')
                            ax.set_ylabel(config['ylabel'], fontsize=9)
                            ax.grid(True, alpha=0.3)
                            
                            if config['ylim']:
                                ax.set_ylim(config['ylim'])
                            
                            # Adicionar estatísticas
                            mean_val = valid_data.mean()
                            max_val = valid_data.max()
                            min_val = valid_data.min()
                            
                            # Formatação baseada na unidade
                            if 'GB' in config['ylabel']:
                                stats_text = f'Média: {mean_val:.1f}GB\nMáx: {max_val:.1f}GB\nMín: {min_val:.1f}GB'
                            elif '%' in config['ylabel']:
                                stats_text = f'Média: {mean_val:.1f}%\nMáx: {max_val:.1f}%\nMín: {min_val:.1f}%'
                                
                                # Adicionar comparação com dados de otimização para CPU
                                if 'CPU' in metric_name and opt_info and opt_info['current_cpu_util'] > 0:
                                    opt_cpu = opt_info['current_cpu_util']
                                    stats_text += f'\nOtim: {opt_cpu:.1f}%'
                            else:
                                stats_text = f'Média: {mean_val:.0f}\nMáx: {max_val:.0f}\nMín: {min_val:.0f}'
                            
                            ax.text(0.02, 0.98, stats_text, transform=ax.transAxes,
                                   verticalalignment='top', fontsize=7,
                                   bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
                        else:
                            ax.text(0.5, 0.5, 'Dados não disponíveis\nou todos são nulos', 
                                   transform=ax.transAxes, ha='center', va='center')
                            ax.set_title(f'{metric_name} - Sem Dados')
                    
                    except Exception as e:
                        ax.text(0.5, 0.5, f'Erro ao processar dados\n{str(e)[:50]}...', 
                               transform=ax.transAxes, ha='center', va='center', fontsize=8)
                        ax.set_title(f'{metric_name} - Erro')
                    
                    # Configurar eixo X
                    ax.tick_params(axis='x', rotation=45, labelsize=8)
                    ax.tick_params(axis='y', labelsize=8)
                
                # Esconder eixos não utilizados
                for idx in range(len(available_metrics), len(axes)):
                    axes[idx].set_visible(False)
                
                # Adicionar box com informações de otimização se disponível
                if opt_info:
                    info_text = self.format_optimization_info(opt_info)
                    fig.text(0.02, 0.02, info_text, fontsize=8,
                            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.8),
                            verticalalignment='bottom')
                
                plt.tight_layout()
                plt.subplots_adjust(bottom=0.15 if opt_info else 0.1)  # Espaço para info de otimização
                
                # Salvar gráfico com resolução otimizada para Word
                safe_instance_name = (instance.replace(':', '_')
                                    .replace('/', '_')
                                    .replace(' ', '_')
                                    .replace('\\', '_')
                                    .replace('|', '_'))[:50]  # Limitar tamanho do nome
                
                chart_path = f"{self.base_output_dir}/charts/{safe_instance_name}_metrics.png"
                plt.savefig(chart_path, dpi=200, bbox_inches='tight', facecolor='white')
                plt.close()
                
                chart_paths[instance] = chart_path
                
            except Exception as e:
                print(f"Erro ao processar instância {instance}: {str(e)}")
                continue
        
        print(f"✓ {len(chart_paths)} gráficos individuais criados")
        return chart_paths
    
    def format_optimization_info(self, opt_info):
        """Formata informações de otimização para exibição com motivos detalhados"""
        if not opt_info:
            return ""
        
        info_lines = []
        
        # Status da otimização com motivos detalhados
        finding = opt_info['finding']
        reason_codes = opt_info.get('finding_reason_codes', '')
        
        if finding == 'Optimized':
            info_lines.append("INSTÂNCIA OTIMIZADA - Apta para Reserva")
            if reason_codes:
                formatted_reasons = self.format_finding_reason_codes(reason_codes)
                if formatted_reasons:
                    info_lines.append(f"Observação: {formatted_reasons}")
        elif finding == 'Underprovisioned':
            info_lines.append("SUB-PROVISIONADA - Recomenda upgrade antes da reserva")
            if reason_codes:
                formatted_reasons = self.format_finding_reason_codes(reason_codes)
                if formatted_reasons:
                    info_lines.append(f"Motivos: {formatted_reasons}")
        elif finding == 'Overprovisioned':
            info_lines.append("SUPER-PROVISIONADA - Potencial economia com downgrade")
            if reason_codes:
                formatted_reasons = self.format_finding_reason_codes(reason_codes)
                if formatted_reasons:
                    info_lines.append(f"Motivos: {formatted_reasons}")
        
        # Classe atual vs recomendada
        current_class = opt_info['current_instance_class']
        recommended_class = opt_info['recommended_instance_class']
        if current_class and recommended_class and current_class != recommended_class:
            info_lines.append(f"Atual: {current_class} → Recomendado: {recommended_class}")
        
        # Custos
        current_cost = opt_info['current_ondemand_monthly']
        recommended_cost = opt_info['recommended_ondemand_monthly']
        savings = opt_info['ondemand_monthly_savings']
        
        if current_cost > 0:
            info_lines.append(f"Custo atual: ${current_cost:.2f}/mês")
            if recommended_cost > 0 and recommended_cost != current_cost:
                info_lines.append(f"Custo recomendado: ${recommended_cost:.2f}/mês")
            if savings != 0:
                if savings > 0:
                    info_lines.append(f"Economia: ${savings:.2f}/mês")
                else:
                    info_lines.append(f"Custo adicional: ${abs(savings):.2f}/mês")
        
        # Melhor opção de reserva para classe atual ou recomendada
        use_recommended = finding != 'Optimized' and recommended_cost != current_cost
        best_reservation = self.get_best_reservation_option(opt_info, use_recommended)
        if best_reservation:
            class_label = "recomendada" if use_recommended else "atual"
            info_lines.append(f"Melhor reserva (classe {class_label}): {best_reservation['term']} {best_reservation['type']}")
            info_lines.append(f"Custo anual: ${best_reservation['annual_cost']:.0f}")
        
        return "\n".join(info_lines)
    
    def create_summary_dashboards(self):
        """Cria dashboards resumo - um gráfico por página para todas as instâncias"""
        print("\nCriando dashboards resumo...")
        
        try:
            # Calcular estatísticas resumo
            summary_stats = []
            
            for i, instance in enumerate(self.instance_names):
                stats = {'Instance': instance}
                
                # Adicionar informações de otimização se disponível
                opt_info = self.get_cost_optimization_info(instance) if self.cost_optimization_data else None
                if opt_info:
                    stats['Finding'] = opt_info['finding']
                    stats['Finding_Reason_Codes'] = opt_info.get('finding_reason_codes', '')  # NOVA COLUNA
                    stats['Current_Class'] = opt_info['current_instance_class']
                    stats['Recommended_Class'] = opt_info['recommended_instance_class']
                    stats['Monthly_Cost'] = opt_info['current_ondemand_monthly']
                    stats['Recommended_Cost'] = opt_info['recommended_ondemand_monthly']
                    stats['Monthly_Savings'] = opt_info['ondemand_monthly_savings']
                    
                    # Melhor opção de reserva
                    use_recommended = opt_info['finding'] != 'Optimized' and opt_info['recommended_ondemand_monthly'] != opt_info['current_ondemand_monthly']
                    best_reservation = self.get_best_reservation_option(opt_info, use_recommended)
                    if best_reservation:
                        stats['Best_Reservation'] = f"{best_reservation['term']} {best_reservation['type']}"
                        stats['Reservation_Annual_Cost'] = best_reservation['annual_cost']
                    else:
                        stats['Best_Reservation'] = 'N/A'
                        stats['Reservation_Annual_Cost'] = 0
                else:
                    stats['Finding'] = 'Unknown'
                    stats['Finding_Reason_Codes'] = ''
                    stats['Current_Class'] = 'Unknown'
                    stats['Recommended_Class'] = ''
                    stats['Monthly_Cost'] = 0
                    stats['Recommended_Cost'] = 0
                    stats['Monthly_Savings'] = 0
                    stats['Best_Reservation'] = 'N/A'
                    stats['Reservation_Annual_Cost'] = 0
                
                for metric_name in self.metrics_data.keys():
                    try:
                        if i < len(self.metrics_data[metric_name].columns):
                            metric_values = self.metrics_data[metric_name].iloc[:, i]
                            valid_data = metric_values.dropna()
                            
                            if len(valid_data) > 0:
                                stats[f'{metric_name}_avg'] = valid_data.mean()
                                stats[f'{metric_name}_max'] = valid_data.max()
                            else:
                                stats[f'{metric_name}_avg'] = 0
                                stats[f'{metric_name}_max'] = 0
                        else:
                            stats[f'{metric_name}_avg'] = 0
                            stats[f'{metric_name}_max'] = 0
                    except:
                        stats[f'{metric_name}_avg'] = 0
                        stats[f'{metric_name}_max'] = 0
                
                summary_stats.append(stats)
            
            stats_df = pd.DataFrame(summary_stats)
            
            # Criar um gráfico por métrica (todas as instâncias)
            dashboard_paths = []
            available_metrics = list(self.metrics_data.keys())
            colors_list = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FECA57', '#FF9FF3']
            
            for idx, metric_name in enumerate(available_metrics):
                print(f"Criando dashboard para {metric_name}...")
                
                # Ajustar tamanho da figura para Word
                fig, ax = plt.subplots(1, 1, figsize=(10, 12))  
                color = colors_list[idx % len(colors_list)]
                
                avg_col = f'{metric_name}_avg'
                if avg_col in stats_df.columns:
                    # Filtrar valores válidos (não zero e não NaN)
                    valid_data = stats_df[stats_df[avg_col] > 0].copy()
                    
                    if len(valid_data) > 0:
                        # Ordernar por valor para melhor visualização
                        sorted_data = valid_data.sort_values(avg_col, ascending=True)
                        
                        # Criar cores baseadas no status de otimização
                        colors = []
                        for _, row in sorted_data.iterrows():
                            finding = row.get('Finding', 'Unknown')
                            if finding == 'Optimized':
                                colors.append('#28a745')  # Verde
                            elif finding == 'Underprovisioned':
                                colors.append('#ffc107')  # Amarelo
                            elif finding == 'Overprovisioned':
                                colors.append('#fd7e14')  # Laranja
                            else:
                                colors.append('#6c757d')  # Cinza
                        
                        # Truncar nomes das instâncias para visualização (sem emojis)
                        labels = []
                        for _, row in sorted_data.iterrows():
                            name = row['Instance']
                            finding = row.get('Finding', '')
                            status_label = ''
                            if finding == 'Optimized':
                                status_label = '[OPT] '
                            elif finding == 'Underprovisioned':
                                status_label = '[SUB] '
                            elif finding == 'Overprovisioned':
                                status_label = '[SUPER] '
                            
                            label = f"{status_label}{name[:20]}{'...' if len(name) > 20 else ''}"
                            labels.append(label)
                        
                        bars = ax.barh(range(len(sorted_data)), sorted_data[avg_col], color=colors, alpha=0.7)
                        ax.set_yticks(range(len(sorted_data)))
                        ax.set_yticklabels(labels, fontsize=8)
                        ax.set_title(f'Resumo Geral - {metric_name} (Todas as Instâncias)', fontsize=12, fontweight='bold', pad=15)
                        
                        # Adicionar valores nas barras com formatação adequada
                        for i, bar in enumerate(bars):
                            width = bar.get_width()
                            if 'Memory' in metric_name or 'Storage' in metric_name:
                                ax.text(width, bar.get_y() + bar.get_height()/2, 
                                       f'{width:.1f}GB', ha='left', va='center', fontsize=7)
                                ax.set_xlabel('GB', fontsize=10)
                            elif 'CPU' in metric_name:
                                ax.text(width, bar.get_y() + bar.get_height()/2, 
                                       f'{width:.1f}%', ha='left', va='center', fontsize=7)
                                ax.set_xlabel('Percentage (%)', fontsize=10)
                            else:
                                ax.text(width, bar.get_y() + bar.get_height()/2, 
                                       f'{width:.0f}', ha='left', va='center', fontsize=7)
                                ax.set_xlabel('Value', fontsize=10)
                        
                        # Adicionar grid para melhor leitura
                        ax.grid(True, alpha=0.3, axis='x')
                        
                        # Adicionar estatísticas gerais no gráfico
                        total_instances = len(sorted_data)
                        avg_value = sorted_data[avg_col].mean()
                        max_value = sorted_data[avg_col].max()
                        min_value = sorted_data[avg_col].min()
                        
                        # Estatísticas de otimização
                        optimized_count = len(sorted_data[sorted_data['Finding'] == 'Optimized'])
                        under_count = len(sorted_data[sorted_data['Finding'] == 'Underprovisioned'])
                        over_count = len(sorted_data[sorted_data['Finding'] == 'Overprovisioned'])
                        
                        if 'Memory' in metric_name or 'Storage' in metric_name:
                            stats_text = f'Total: {total_instances}\nMédia: {avg_value:.1f}GB\nMáx: {max_value:.1f}GB\nMín: {min_value:.1f}GB'
                        elif 'CPU' in metric_name:
                            stats_text = f'Total: {total_instances}\nMédia: {avg_value:.1f}%\nMáx: {max_value:.1f}%\nMín: {min_value:.1f}%'
                        else:
                            stats_text = f'Total: {total_instances}\nMédia: {avg_value:.0f}\nMáx: {max_value:.0f}\nMín: {min_value:.0f}'
                        
                        stats_text += f'\n\nOtimização:\nOtimizadas: {optimized_count}\nSub-prov.: {under_count}\nSuper-prov.: {over_count}'
                        
                        ax.text(0.98, 0.02, stats_text, transform=ax.transAxes,
                               verticalalignment='bottom', horizontalalignment='right', fontsize=9,
                               bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
                    else:
                        ax.text(0.5, 0.5, f'Nenhum dado válido para {metric_name}', 
                               transform=ax.transAxes, ha='center', va='center', fontsize=14)
                        ax.set_title(f'{metric_name} - Sem Dados', fontsize=14, fontweight='bold')
                
                plt.tight_layout()
                dashboard_path = f"{self.base_output_dir}/dashboard_{metric_name.lower().replace(' ', '_')}.png"
                plt.savefig(dashboard_path, dpi=200, bbox_inches='tight', facecolor='white')
                plt.close()
                
                dashboard_paths.append(dashboard_path)
            
            print(f"✓ {len(dashboard_paths)} dashboards de resumo criados")
            return dashboard_paths, stats_df
            
        except Exception as e:
            print(f"✗ Erro ao criar dashboards resumo: {str(e)}")
            return [], pd.DataFrame()
    
    def add_page_break(self, doc):
        """Adiciona quebra de página simples"""
        from docx.enum.text import WD_BREAK
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def generate_word_report(self, chart_paths, dashboard_paths, stats_df):
        """Gera relatório Word com todos os gráficos"""
        
        # Usar template se fornecido, senão criar documento em branco
        if self.template_path and os.path.exists(self.template_path):
            print(f"Usando template: {self.template_path}")
            doc = Document(self.template_path)
            # Adicionar quebra de página após o template
            self.add_page_break(doc)
        else:
            print("Criando documento em branco (template não encontrado ou não fornecido)")
            doc = Document()
        
        word_path = f"{self.base_output_dir}/RDS_Performance_Report.docx"
        print(f"\nGerando relatório Word: {word_path}")
        
        try:
            # Título principal (apenas se não usar template)
            if not (self.template_path and os.path.exists(self.template_path)):
                title = doc.add_heading('Relatório de Performance - Instâncias RDS', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informações gerais
            info_heading = doc.add_heading('Informações Gerais', level=1)
            
            period_start = self.dates.min().strftime('%d/%m/%Y %H:%M') if self.dates is not None else 'N/A'
            period_end = self.dates.max().strftime('%d/%m/%Y %H:%M') if self.dates is not None else 'N/A'
            
            info_para = doc.add_paragraph()
            info_para.add_run('Período de Análise: ').bold = True
            info_para.add_run(f'{period_start} até {period_end}\n')
            info_para.add_run('Total de Instâncias: ').bold = True
            info_para.add_run(f'{len(self.instance_names)}\n')
            info_para.add_run('Métricas Analisadas: ').bold = True
            info_para.add_run(f'{", ".join(self.metrics_data.keys())}\n')
            info_para.add_run('Data de Geração: ').bold = True
            info_para.add_run(f'{datetime.now().strftime("%d/%m/%Y %H:%M")}\n')
            info_para.add_run('Total de Pontos de Dados: ').bold = True
            info_para.add_run(f'{len(self.dates) if self.dates is not None else 0} por instância')
            
            # Resumo de otimização se disponível
            if self.cost_optimization_data:
                doc.add_paragraph()
                opt_para = doc.add_paragraph()
                opt_para.add_run('Análise de Otimização de Custos:').bold = True
                
                # Contar status de otimização
                optimized_count = sum(1 for data in self.cost_optimization_data.values() 
                                    if data['finding'] == 'Optimized' and data['status'] == 'available')
                under_count = sum(1 for data in self.cost_optimization_data.values() 
                                if data['finding'] == 'Underprovisioned' and data['status'] == 'available')
                over_count = sum(1 for data in self.cost_optimization_data.values() 
                               if data['finding'] == 'Overprovisioned' and data['status'] == 'available')
                
                opt_para.add_run(f'\nInstâncias Otimizadas (aptas para reserva): {optimized_count}')
                opt_para.add_run(f'\nInstâncias Sub-provisionadas: {under_count}')
                opt_para.add_run(f'\nInstâncias Super-provisionadas: {over_count}')
            
            # Quebra de página
            self.add_page_break(doc)
            
            # Dashboards resumo
            doc.add_heading('Resumo Geral por Métrica', level=1)
            
            for i, dashboard_path in enumerate(dashboard_paths):
                if i > 0:  # Quebra de página para cada dashboard exceto o primeiro
                    self.add_page_break(doc)
                
                if os.path.exists(dashboard_path):
                    # Adicionar imagem do dashboard
                    doc.add_picture(dashboard_path, width=Inches(6.5))
                    
                    # Centralizar a imagem
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"Dashboard não encontrado: {dashboard_path}")
            
            # Quebra de página
            self.add_page_break(doc)
            
            # Tabela de estatísticas com otimização E MOTIVOS
            if not stats_df.empty:
                doc.add_heading('Análise de Otimização e Recomendações de Reserva', level=1)
                
                # Criar cabeçalho da tabela incluindo MOTIVOS
                headers = ['Instância', 'Status', 'Motivos', 'Classe Atual', 'Classe Recomendada', 'Custo Mensal', 'Melhor Reserva', 'Custo Anual Reserva']
                for metric in self.metrics_data.keys():
                    if 'CPU' in metric:
                        headers.extend(['CPU Média (%)', 'CPU Máx (%)'])
                    elif 'Storage' in metric or 'FreeStorageSpace' in metric:
                        headers.extend(['Storage Média (GB)', 'Storage Máx (GB)'])
                    elif 'Memory' in metric or 'FreeableMemory' in metric:
                        headers.extend(['Mem Média (GB)', 'Mem Máx (GB)'])
                    elif 'Connection' in metric or 'DatabaseConnection' in metric:
                        headers.extend(['Conn Média', 'Conn Máx'])
                    else:
                        short_name = metric[:8] + '...' if len(metric) > 8 else metric
                        headers.extend([f'{short_name} Média', f'{short_name} Máx'])
                
                # Criar tabela
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Adicionar cabeçalhos
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    # Fazer cabeçalho em negrito
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adicionar dados
                for _, row in stats_df.iterrows():
                    row_cells = table.add_row().cells
                    
                    # Nome da instância
                    instance_name = row['Instance'][:25] + '...' if len(row['Instance']) > 25 else row['Instance']
                    row_cells[0].text = instance_name
                    
                    # Informações de otimização
                    finding = row.get('Finding', 'Unknown')
                    if finding == 'Optimized':
                        row_cells[1].text = 'OTIMIZADA'
                    elif finding == 'Underprovisioned':
                        row_cells[1].text = 'SUB-PROV'
                    elif finding == 'Overprovisioned':
                        row_cells[1].text = 'SUPER-PROV'
                    else:
                        row_cells[1].text = 'DESCONHECIDO'
                    
                    # NOVA COLUNA: Motivos traduzidos
                    reason_codes = row.get('Finding_Reason_Codes', '')
                    formatted_reasons = self.format_finding_reason_codes(reason_codes)
                    row_cells[2].text = formatted_reasons[:50] + '...' if len(formatted_reasons) > 50 else formatted_reasons
                    
                    row_cells[3].text = row.get('Current_Class', 'N/A')[:15]
                    row_cells[4].text = row.get('Recommended_Class', 'N/A')[:15]
                    
                    monthly_cost = row.get('Monthly_Cost', 0)
                    row_cells[5].text = f"${monthly_cost:.0f}" if monthly_cost > 0 else 'N/A'
                    
                    # Melhor modalidade de reserva
                    best_reservation = row.get('Best_Reservation', 'N/A')
                    row_cells[6].text = best_reservation
                    
                    # Custo anual da reserva
                    reservation_cost = row.get('Reservation_Annual_Cost', 0)
                    row_cells[7].text = f"${reservation_cost:.0f}" if reservation_cost > 0 else 'N/A'
                    
                    col_idx = 8  # Começar após a nova coluna "Motivos"
                    for metric in self.metrics_data.keys():
                        avg_col = f'{metric}_avg'
                        max_col = f'{metric}_max'
                        if avg_col in row and max_col in row:
                            avg_val = row[avg_col] if not pd.isna(row[avg_col]) else 0
                            max_val = row[max_col] if not pd.isna(row[max_col]) else 0
                            
                            # Formatação baseada na métrica
                            if 'Memory' in metric or 'Storage' in metric:
                                row_cells[col_idx].text = f"{avg_val:.1f}"
                                row_cells[col_idx + 1].text = f"{max_val:.1f}"
                            elif 'CPU' in metric:
                                row_cells[col_idx].text = f"{avg_val:.1f}"
                                row_cells[col_idx + 1].text = f"{max_val:.1f}"
                            else:
                                row_cells[col_idx].text = f"{avg_val:.0f}"
                                row_cells[col_idx + 1].text = f"{max_val:.0f}"
                        else:
                            row_cells[col_idx].text = 'N/A'
                            row_cells[col_idx + 1].text = 'N/A'
                        
                        # Centralizar texto nas células
                        row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        col_idx += 2
                
                # Ajustar largura das colunas
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(0.7)
            
            # Quebra de página
            self.add_page_break(doc)
            
            # Gráficos individuais
            doc.add_heading('Métricas Detalhadas por Instância', level=1)
            
            for i, (instance, chart_path) in enumerate(chart_paths.items()):
                # Quebra de página para cada instância
                if i > 0:
                    self.add_page_break(doc)
                
                # Nome da instância com status de otimização (sem emojis)
                display_name = instance[:80] + '...' if len(instance) > 80 else instance
                opt_info = self.get_cost_optimization_info(instance) if self.cost_optimization_data else None
                
                if opt_info:
                    finding = opt_info['finding']
                    if finding == 'Optimized':
                        display_name += ' [OTIMIZADA]'
                    elif finding == 'Underprovisioned':
                        display_name += ' [SUB-PROVISIONADA]'
                    elif finding == 'Overprovisioned':
                        display_name += ' [SUPER-PROVISIONADA]'
                
                doc.add_heading(f'Instância: {display_name}', level=2)
                
                # Adicionar informações de otimização detalhadas COM MOTIVOS
                if opt_info:
                    opt_detail = doc.add_paragraph()
                    
                    finding = opt_info['finding']
                    reason_codes = opt_info.get('finding_reason_codes', '')
                    current_class = opt_info['current_instance_class']
                    recommended_class = opt_info['recommended_instance_class']
                    current_cost = opt_info['current_ondemand_monthly']
                    recommended_cost = opt_info['recommended_ondemand_monthly']
                    savings = opt_info['ondemand_monthly_savings']
                    
                    # Determinar se deve usar classe recomendada
                    use_recommended = finding != 'Optimized' and recommended_cost != current_cost and recommended_class
                    
                    if finding == 'Optimized':
                        opt_detail.add_run('Status: ').bold = True
                        opt_detail.add_run('OTIMIZADA - Apta para reserva\n')
                        if reason_codes:
                            formatted_reasons = self.format_finding_reason_codes(reason_codes)
                            if formatted_reasons:
                                opt_detail.add_run('Observação: ').bold = True
                                opt_detail.add_run(f'{formatted_reasons}\n')
                    elif finding == 'Underprovisioned':
                        opt_detail.add_run('Status: ').bold = True
                        opt_detail.add_run('SUB-PROVISIONADA - Considerar upgrade antes da reserva\n')
                        if reason_codes:
                            formatted_reasons = self.format_finding_reason_codes(reason_codes)
                            if formatted_reasons:
                                opt_detail.add_run('Motivos: ').bold = True
                                opt_detail.add_run(f'{formatted_reasons}\n')
                    elif finding == 'Overprovisioned':
                        opt_detail.add_run('Status: ').bold = True
                        opt_detail.add_run('SUPER-PROVISIONADA - Oportunidade de economia\n')
                        if reason_codes:
                            formatted_reasons = self.format_finding_reason_codes(reason_codes)
                            if formatted_reasons:
                                opt_detail.add_run('Motivos: ').bold = True
                                opt_detail.add_run(f'{formatted_reasons}\n')
                    
                    if current_class:
                        opt_detail.add_run('Classe Atual: ').bold = True
                        opt_detail.add_run(f'{current_class}\n')
                    
                    if recommended_class and recommended_class != current_class:
                        opt_detail.add_run('Classe Recomendada: ').bold = True
                        opt_detail.add_run(f'{recommended_class}\n')
                    
                    # Custos On-Demand
                    if current_cost > 0:
                        opt_detail.add_run('Custo On-Demand (classe atual): ').bold = True
                        opt_detail.add_run(f'${current_cost:.2f}/mês\n')
                    
                    if use_recommended and recommended_cost > 0:
                        opt_detail.add_run('Custo On-Demand (classe recomendada): ').bold = True
                        opt_detail.add_run(f'${recommended_cost:.2f}/mês\n')
                    
                    # Opções de reserva apenas para a classe a ser utilizada (atual ou recomendada)
                    if use_recommended:
                        opt_detail.add_run('\nOpções de Reserva (classe recomendada):\n').bold = True
                        recommended_options = self.get_reservation_options(opt_info, True)
                        for option in recommended_options:
                            opt_detail.add_run(f'• {option["term"]} {option["type"]}: {option["display"]}\n')
                    else:
                        # Instância otimizada - mostrar opções para classe atual
                        opt_detail.add_run('\nOpções de Reserva:\n').bold = True
                        current_options = self.get_reservation_options(opt_info, False)
                        for option in current_options:
                            opt_detail.add_run(f'• {option["term"]} {option["type"]}: {option["display"]}\n')
                    
                    # Melhor recomendação
                    best_reservation = self.get_best_reservation_option(opt_info, use_recommended)
                    if best_reservation:
                        opt_detail.add_run('\nMelhor Opção: ').bold = True
                        opt_detail.add_run(f'{best_reservation["term"]} {best_reservation["type"]}\n')
                        
                        if best_reservation['upfront_cost'] > 0:
                            opt_detail.add_run(f'• Pagamento inicial: ${best_reservation["upfront_cost"]:.0f}\n')
                        
                        if best_reservation['monthly_cost'] > 0:
                            opt_detail.add_run(f'• Pagamento mensal: ${best_reservation["monthly_cost"]:.2f}\n')
                        
                        opt_detail.add_run(f'• Custo total anual: ${best_reservation["annual_cost"]:.0f}\n')
                        
                        # Economia vs On-Demand
                        comparison_cost = recommended_cost if use_recommended else current_cost
                        if comparison_cost > 0:
                            monthly_savings = comparison_cost - best_reservation['annual_cost']/12
                            annual_savings = monthly_savings * 12
                            savings_percentage = (monthly_savings / comparison_cost) * 100
                            
                            if monthly_savings > 0:
                                opt_detail.add_run('• Economia vs On-Demand: ').bold = True
                                opt_detail.add_run(f'${monthly_savings:.2f}/mês (${annual_savings:.0f}/ano - {savings_percentage:.1f}%)')
                            elif monthly_savings < 0:
                                opt_detail.add_run('• Custo adicional vs On-Demand: ').bold = True
                                opt_detail.add_run(f'${abs(monthly_savings):.2f}/mês (${abs(annual_savings):.0f}/ano)')
                
                if os.path.exists(chart_path):
                    # Adicionar gráfico da instância
                    doc.add_picture(chart_path, width=Inches(6.5))
                    
                    # Centralizar a imagem
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"Gráfico não encontrado: {chart_path}")
            
            # Salvar documento
            doc.save(word_path)
            print(f"✓ Relatório Word gerado com sucesso: {word_path}")
            return word_path
            
        except Exception as e:
            print(f"✗ Erro ao gerar Word: {str(e)}")
            return None
    
    def generate_complete_report(self, csv_files):
        """Função principal para gerar relatório completo"""
        print("=" * 80)
        print("     INICIANDO GERAÇÃO DO RELATÓRIO RDS PERFORMANCE + OTIMIZAÇÃO")
        print("=" * 80)
        
        # 1. Carregar dados de otimização
        print("\n1. CARREGANDO DADOS DE OTIMIZAÇÃO...")
        self.load_cost_optimization_data()
        
        # 2. Carregar dados de métricas
        print("\n2. CARREGANDO DADOS DE MÉTRICAS...")
        self.load_metric_data(csv_files)
        
        if not self.metrics_data:
            print("❌ ERRO: Nenhum dado de métrica foi carregado!")
            print("Verifique se os arquivos CSV existem e estão no formato correto.")
            return None
        
        print(f"✅ Dados carregados com sucesso!")
        print(f"   📊 Métricas: {list(self.metrics_data.keys())}")
        print(f"   🖥️  Instâncias: {len(self.instance_names)}")
        print(f"   📅 Período: {len(self.dates)} pontos de dados")
        
        if self.cost_optimization_data:
            optimized = sum(1 for d in self.cost_optimization_data.values() if d['finding'] == 'Optimized')
            print(f"   💰 Otimização: {len(self.cost_optimization_data)} instâncias analisadas ({optimized} otimizadas)")
        
        # 3. Criar gráficos individuais
        print("\n3. GERANDO GRÁFICOS INDIVIDUAIS...")
        chart_paths = self.create_individual_charts()
        
        if not chart_paths:
            print("❌ ERRO: Nenhum gráfico foi criado!")
            return None
        
        # 4. Criar dashboards resumo
        print("\n4. GERANDO DASHBOARDS RESUMO...")
        dashboard_paths, stats_df = self.create_summary_dashboards()
        
        # 5. Gerar relatório Word
        print("\n5. GERANDO RELATÓRIO WORD...")
        word_path = self.generate_word_report(chart_paths, dashboard_paths, stats_df)
        
        if word_path:
            print("\n" + "=" * 80)
            print("           RELATÓRIO COMPLETO GERADO COM SUCESSO!")
            print("=" * 80)
            print(f"📄 Arquivo Word: {word_path}")
            print(f"📁 Pasta gráficos: {self.base_output_dir}/charts/")
            print(f"📊 Total de gráficos individuais: {len(chart_paths)}")
            print(f"📊 Total de dashboards resumo: {len(dashboard_paths)}")
            
            # Resumo de otimização
            if self.cost_optimization_data:
                print(f"\n💰 RESUMO DE OTIMIZAÇÃO:")
                optimized = sum(1 for d in self.cost_optimization_data.values() if d['finding'] == 'Optimized' and d['status'] == 'available')
                under = sum(1 for d in self.cost_optimization_data.values() if d['finding'] == 'Underprovisioned' and d['status'] == 'available')
                over = sum(1 for d in self.cost_optimization_data.values() if d['finding'] == 'Overprovisioned' and d['status'] == 'available')
                
                print(f"   Instâncias otimizadas (aptas para reserva): {optimized}")
                print(f"   Instâncias sub-provisionadas: {under}")
                print(f"   Instâncias super-provisionadas: {over}")
                
                # Calcular economia total potencial com reservas
                total_current_monthly = sum(d['current_ondemand_monthly'] for d in self.cost_optimization_data.values() if d['current_ondemand_monthly'] > 0)
                total_reservation_annual = 0
                for d in self.cost_optimization_data.values():
                    use_recommended = d['finding'] != 'Optimized' and d['recommended_ondemand_monthly'] != d['current_ondemand_monthly']
                    best_res = self.get_best_reservation_option(d, use_recommended)
                    if best_res:
                        total_reservation_annual += best_res['annual_cost']
                
                if total_current_monthly > 0 and total_reservation_annual > 0:
                    monthly_savings = total_current_monthly - (total_reservation_annual / 12)
                    annual_savings = monthly_savings * 12
                    if monthly_savings > 0:
                        print(f"   💰 Economia potencial com reservas: ${monthly_savings:.2f}/mês (${annual_savings:.0f}/ano)")
            
            return word_path
        else:
            print("\n❌ ERRO: Falha ao gerar o relatório Word!")
            return None

def main():
    """Função principal - Configure aqui os seus arquivos CSV e template"""
    
    print("RDS Performance Report Generator - Versão com Motivos Detalhados")
    print("Desenvolvido para análise de métricas CloudWatch + Recomendações Otimizadas")
    print()
    
     # ============================================
    # CONFIGURAÇÃO - EDITE AQUI OS SEUS ARQUIVOS
    # ============================================
    
    csv_files = {
        'CPU': 'cpu_utilization.csv',
        'FreeStorageSpace': 'free_storage_space.csv',
        'FreeableMemory': 'freeable_memory.csv',
        'DatabaseConnections': 'database_connections.csv',
    }
    
    # ============================================
    # CONFIGURAÇÃO DO TEMPLATE E OTIMIZAÇÃO
    # ============================================
    template_path = None  # Caminho para seu template Word
    cost_optimization_csv = 'rds_cost_optimization.csv'  # CSV de otimização
    # Se não tiver template, deixe como None: template_path = None
    # Se não tiver CSV de otimização, deixe como None: cost_optimization_csv = None
    
    # Verificar se pelo menos um arquivo existe
    existing_files = {k: v for k, v in csv_files.items() if os.path.exists(v)}
    
    if not existing_files:
        print("❌ ERRO: Nenhum arquivo CSV encontrado!")
        print("\nArquivos procurados:")
        for metric, path in csv_files.items():
            status = "✅ ENCONTRADO" if os.path.exists(path) else "❌ NÃO ENCONTRADO"
            print(f"   {metric}: {path} - {status}")
        print("\nPor favor, verifique os caminhos dos arquivos na função main().")
        return
    
    print("Arquivos encontrados:")
    for metric, path in existing_files.items():
        print(f"   ✅ {metric}: {path}")
    
    if len(existing_files) < len(csv_files):
        print(f"\nAviso: {len(csv_files) - len(existing_files)} arquivo(s) não encontrado(s), mas continuando com os disponíveis...")
    
    # Verificar template
    if template_path and os.path.exists(template_path):
        print(f"\n📋 Template encontrado: {template_path}")
    elif template_path:
        print(f"\n⚠️  Template não encontrado: {template_path}")
        print("Continuando sem template (documento em branco)...")
        template_path = None
    else:
        print(f"\n📄 Nenhum template configurado - usando documento em branco")
    
    # Verificar CSV de otimização
    if cost_optimization_csv and os.path.exists(cost_optimization_csv):
        print(f"\n💰 CSV de otimização encontrado: {cost_optimization_csv}")
    elif cost_optimization_csv:
        print(f"\n⚠️  CSV de otimização não encontrado: {cost_optimization_csv}")
        print("Continuando sem análise de otimização...")
        cost_optimization_csv = None
    else:
        print(f"\n📊 Nenhum CSV de otimização configurado")
    
    # Criar instância do gerador
    print(f"\nInicializando gerador de relatórios Word com motivos detalhados...")
    report_generator = RDSReportGeneratorWord('rds_performance_report', template_path, cost_optimization_csv)
    
    # Gerar relatório completo
    try:
        word_path = report_generator.generate_complete_report(existing_files)
        
        if word_path:
            print(f"\n🎉 SUCESSO! Relatório disponível em: {word_path}")
            
            # Tentar abrir o arquivo automaticamente (opcional)
            try:
                import webbrowser
                webbrowser.open(word_path)
                print("📖 Abrindo relatório automaticamente...")
            except:
                print("💡 Abra manualmente o arquivo Word para visualizar o relatório.")
        else:
            print("\n❌ Falha na geração do relatório. Verifique os logs acima para detalhes.")
            
    except Exception as e:
        print(f"\n❌ ERRO CRÍTICO: {str(e)}")
        print("Verifique se todas as dependências estão instaladas:")
        print("pip install pandas matplotlib seaborn python-docx")

if __name__ == "__main__":
    main()
